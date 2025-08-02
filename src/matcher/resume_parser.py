"""
Resume Parser for extracting information from PDF and DOCX files
"""

import re
from typing import Dict, List, Optional
from pathlib import Path

import PyPDF2
import pdfplumber
from docx import Document
import spacy
from spacy.matcher import Matcher

from src.utils.logger import setup_logger


class ResumeParser:
    """Parse resumes and extract relevant information"""
    
    def __init__(self):
        """Initialize the parser"""
        self.logger = setup_logger(__name__)
        
        # Load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            self.logger.warning("spaCy model not found")
            self.nlp = None
            
        self.matcher = Matcher(self.nlp.vocab) if self.nlp else None
        
    def parse_resume(self, file_path: str) -> Dict:
        """Parse resume from file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return {}
            
        # Extract text based on file type
        text = ""
        if file_path.suffix.lower() == '.pdf':
            text = self._extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = self._extract_text_from_docx(str(file_path))
        else:
            self.logger.error(f"Unsupported file format: {file_path.suffix}")
            return {}
            
        if not text:
            self.logger.error(f"No text extracted from: {file_path}")
            return {}
            
        # Parse information
        resume_data = {
            'text': text,
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'skills': self._extract_skills(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text),
            'linkedin': self._extract_linkedin(text),
            'github': self._extract_github(text)
        }
        
        return resume_data
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
                        
            except Exception as e:
                self.logger.error(f"Error extracting PDF text: {e}")
                
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            
        return text
    
    def _extract_name(self, text: str) -> str:
        """Extract name from resume text"""
        # Usually the name is in the first few lines
        lines = text.strip().split('\n')[:5]
        
        for line in lines:
            line = line.strip()
            # Skip empty lines and common headers
            if not line or len(line) < 3:
                continue
                
            # Skip lines with special characters or numbers
            if re.search(r'[@#$%&*()0-9]', line):
                continue
                
            # Check if line contains mostly alphabets and spaces
            if re.match(r'^[A-Za-z\s\.]+$', line) and len(line.split()) <= 4:
                return line
                
        return "Unknown"
    
    def _extract_email(self, text: str) -> str:
        """Extract email from resume text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number from resume text"""
        # Indian phone number patterns
        phone_patterns = [
            r'[\+]?91[-.\s]?[6-9]\d{9}',
            r'[6-9]\d{9}',
            r'[\+]?[1-9]\d{1,14}',
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
                
        return ""
    
    def _extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn URL from resume text"""
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else ""
    
    def _extract_github(self, text: str) -> str:
        """Extract GitHub URL from resume text"""
        github_pattern = r'github\.com/[\w-]+'
        matches = re.findall(github_pattern, text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        skills = set()
        
        # Programming languages
        prog_languages = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go',
            'rust', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'shell',
            'bash', 'powershell', 'sql', 'html', 'css', 'sass', 'less'
        ]
        
        # Frameworks and libraries
        frameworks = [
            'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'spring', 'spring boot',
            'nodejs', 'node.js', 'express', 'rails', 'laravel', 'asp.net', 'jquery',
            'bootstrap', 'tailwind', 'material-ui', 'tensorflow', 'pytorch', 'keras',
            'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'opencv'
        ]
        
        # Tools and platforms
        tools = [
            'git', 'github', 'gitlab', 'bitbucket', 'docker', 'kubernetes', 'jenkins',
            'aws', 'azure', 'gcp', 'heroku', 'firebase', 'mongodb', 'mysql', 'postgresql',
            'redis', 'elasticsearch', 'kafka', 'rabbitmq', 'nginx', 'apache', 'linux',
            'windows', 'macos', 'android', 'ios', 'react native', 'flutter', 'unity'
        ]
        
        # Data science and ML
        ds_ml = [
            'machine learning', 'deep learning', 'artificial intelligence', 'data science',
            'data analysis', 'data mining', 'nlp', 'natural language processing',
            'computer vision', 'neural networks', 'reinforcement learning', 'big data',
            'hadoop', 'spark', 'tableau', 'power bi', 'excel', 'statistics'
        ]
        
        # Combine all skills
        all_skills = prog_languages + frameworks + tools + ds_ml
        
        # Search for skills in text
        text_lower = text.lower()
        for skill in all_skills:
            if skill in text_lower:
                skills.add(skill)
                
        # Also extract skills using regex patterns
        skill_patterns = [
            r'\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\b',  # CamelCase
            r'\b([A-Z]+(?:\.[A-Z]+)*)\b',  # Acronyms like AWS, GCP
        ]
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if 2 < len(match) < 20:  # Reasonable skill name length
                    skills.add(match.lower())
                    
        return list(skills)
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information"""
        education = []
        
        # Common degree patterns
        degree_patterns = [
            r'(B\.?Tech|B\.?E\.?|Bachelor)',
            r'(M\.?Tech|M\.?E\.?|Master)',
            r'(B\.?Sc|BSc|Bachelor of Science)',
            r'(M\.?Sc|MSc|Master of Science)',
            r'(MBA|Master of Business)',
            r'(B\.?Com|BCom|Bachelor of Commerce)',
            r'(M\.?Com|MCom|Master of Commerce)',
            r'(Ph\.?D|PhD|Doctorate)',
            r'(B\.?C\.?A|BCA)',
            r'(M\.?C\.?A|MCA)'
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            for pattern in degree_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    education_entry = {
                        'degree': line.strip(),
                        'details': ''
                    }
                    
                    # Try to get next few lines for more details
                    if i + 1 < len(lines):
                        education_entry['details'] = lines[i + 1].strip()
                        
                    education.append(education_entry)
                    break
                    
        return education
    
    def _extract_experience(self, text: str) -> List[str]:
        """Extract work experience sections"""
        experience = []
        
        # Common section headers
        experience_headers = [
            'work experience', 'professional experience', 'employment history',
            'experience', 'career summary', 'professional summary'
        ]
        
        lines = text.split('\n')
        in_experience_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if we're entering experience section
            for header in experience_headers:
                if header in line_lower:
                    in_experience_section = True
                    break
                    
            # Check if we're leaving experience section
            if in_experience_section and any(keyword in line_lower for keyword in 
                ['education', 'skills', 'projects', 'certifications', 'awards']):
                break
                
            # Collect experience lines
            if in_experience_section and line.strip():
                experience.append(line.strip())
                
        return experience